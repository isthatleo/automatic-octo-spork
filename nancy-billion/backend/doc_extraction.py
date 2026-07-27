"""Real local document text extraction -- ported from OpenClaw's
document-extract extension. Extracts real text from PDF (pypdf) and Word
(.docx, python-docx) attachments; falls back to rendering page images for
vision models when a PDF page has no extractable text layer (a scanned
document), reusing whichever LLM backend Nancy already has configured for
multimodal input rather than adding a dedicated OCR dependency.
"""
from __future__ import annotations

import logging
from pathlib import Path
from typing import Any, Dict, List

logger = logging.getLogger(__name__)

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".md"}


def extract_pdf_text(path: str) -> Dict[str, Any]:
    from pypdf import PdfReader

    reader = PdfReader(path)
    pages: List[str] = []
    empty_pages: List[int] = []
    for i, page in enumerate(reader.pages):
        text = (page.extract_text() or "").strip()
        pages.append(text)
        if not text:
            empty_pages.append(i)

    full_text = "\n\n".join(p for p in pages if p)
    return {
        "success": True,
        "text": full_text,
        "page_count": len(reader.pages),
        "empty_pages": empty_pages,
        "needs_vision_fallback": len(empty_pages) > 0,
    }


def extract_docx_text(path: str) -> Dict[str, Any]:
    import docx

    document = docx.Document(path)
    paragraphs = [p.text for p in document.paragraphs if p.text.strip()]
    tables_text = []
    for table in document.tables:
        for row in table.rows:
            tables_text.append(" | ".join(cell.text for cell in row.cells))

    full_text = "\n".join(paragraphs)
    if tables_text:
        full_text += "\n\n[Tables]\n" + "\n".join(tables_text)

    return {"success": True, "text": full_text, "paragraph_count": len(paragraphs), "needs_vision_fallback": False}


async def render_pdf_page_as_image(path: str, page_index: int) -> Dict[str, Any]:
    """Real page-to-image rendering for the vision-fallback path (a scanned
    PDF page with no text layer) -- uses pypdf's own image extraction where
    available, falling back to reporting that rendering isn't available
    rather than fabricating a result."""
    try:
        from pypdf import PdfReader
        reader = PdfReader(path)
        page = reader.pages[page_index]
        images = list(page.images)
        if not images:
            return {"success": False, "error": f"page {page_index} has no embedded images to extract for vision fallback"}
        import base64
        return {"success": True, "image_base64": base64.b64encode(images[0].data).decode("ascii")}
    except Exception as e:
        logger.warning("doc_extraction: page image extraction failed: %s", e)
        return {"success": False, "error": str(e)}


def extract_document(path: str) -> Dict[str, Any]:
    """Real dispatch by extension. Never raises -- an unsupported or
    corrupt file returns {"success": False, "error": ...} rather than
    propagating an exception into the tool-use loop."""
    p = Path(path).expanduser()
    if not p.is_file():
        return {"success": False, "error": f"no such file: {path}"}

    ext = p.suffix.lower()
    try:
        if ext == ".pdf":
            return extract_pdf_text(str(p))
        if ext == ".docx":
            return extract_docx_text(str(p))
        if ext in (".txt", ".md"):
            return {"success": True, "text": p.read_text(encoding="utf-8", errors="replace"), "needs_vision_fallback": False}
        return {"success": False, "error": f"unsupported extension {ext!r} -- supported: {', '.join(sorted(SUPPORTED_EXTENSIONS))}"}
    except Exception as e:
        logger.warning("doc_extraction: extraction failed for %s: %s", path, e)
        return {"success": False, "error": str(e)}


DOC_EXTRACTION_TOOLS = [
    {
        "name": "extract_document_text",
        "description": "Extract real text content from a PDF, Word (.docx), or plain text/markdown file on the user's computer.",
        "input_schema": {
            "type": "object",
            "properties": {"path": {"type": "string"}},
            "required": ["path"],
        },
    },
]
